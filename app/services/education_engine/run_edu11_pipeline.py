import os
import json
import traceback
import pytesseract

from docx import Document
from pdf2image import convert_from_path
from multiprocessing import Pool, cpu_count

# --- IMPORT PARSERS ---
from app.services.education_engine.education_parser import extract_education
from app.services.education_engine.certification_parser import extract_certifications


# === CONFIG ===
RAW_DIR = "data/raw"
OUTPUT_DIR = "data/processed/output_11"

POPPLER_PATH = r"D:\Poppler-25.12.0\Library\bin"
TESSERACT_PATH = r"D:\tesseract\tesseract.exe"

pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

os.makedirs(OUTPUT_DIR, exist_ok=True)


# === LOGGING ===
def log(msg):
    print(f"[LOG] {msg}")


def log_error(file, err):
    print(f"[ERROR] {file}: {err}")


# === FILE VALIDATION ===
def is_valid_file(path):
    return os.path.exists(path) and os.path.isfile(path)


# === DOCX READER ===
def read_docx(path):
    try:
        doc = Document(path)
        text = []

        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text.strip())

        return "\n".join(text)

    except Exception as e:
        log_error(path, f"DOCX read failed: {e}")
        return ""


# === PDF READER (DIRECT) ===
def read_pdf(path):
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(path)
        text = []

        for page in reader.pages:
            content = page.extract_text()
            if content:
                text.append(content)

        return "\n".join(text)

    except Exception as e:
        log_error(path, f"PDF read failed: {e}")
        return ""


# === OCR FALLBACK ===
def ocr_fallback(path):
    try:
        log(f"OCR fallback triggered for {path}")

        images = convert_from_path(path, poppler_path=POPPLER_PATH)
        text = []

        for img in images:
            text.append(pytesseract.image_to_string(img))

        return "\n".join(text)

    except Exception as e:
        log_error(path, f"OCR failed: {e}")
        return ""


# === PROCESS SINGLE FILE ===
def process_file(file):
    try:
        path = os.path.join(RAW_DIR, file)

        if not is_valid_file(path):
            log_error(file, "Invalid file path")
            return None

        log(f"Processing: {file}")

        # --- READ FILE ---
        if file.lower().endswith(".pdf"):
            text = read_pdf(path)

            # fallback if empty
            if not text.strip():
                text = ocr_fallback(path)

        elif file.lower().endswith(".docx"):
            text = read_docx(path)

        elif file.lower().endswith(".txt"):
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()

        else:
            log(f"Skipping unsupported file: {file}")
            return None

        if not text.strip():
            log_error(file, "Empty text after extraction")
            return None

        # === 🔥 EXTRACTION ===
        education = extract_education(text)
        certifications = extract_certifications(text)

        result = {
            "file": file,
            "text_preview": text[:500],
            "education": education,
            "certifications": certifications
        }

        # --- SAVE OUTPUT ---
        output_file = os.path.splitext(file)[0] + ".json"
        output_path = os.path.join(OUTPUT_DIR, output_file)

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=2)

        log(f"Saved: {output_path}")
        return result

    except Exception:
        log_error(file, traceback.format_exc())
        return None


# === RUN PIPELINE ===
def run(debug=True):
    files = [
        f for f in os.listdir(RAW_DIR)
        if f.lower().endswith((".pdf", ".docx", ".txt"))
    ]

    if not files:
        print("No files found.")
        return

    print(f"Processing {len(files)} files...\n")

    # DEBUG MODE (no multiprocessing)
    if debug:
        for f in files:
            result = process_file(f)
            if result:
                print(result)

    # PRODUCTION MODE
    else:
        with Pool(cpu_count() - 1) as p:
            results = p.map(process_file, files)

        for r in results:
            if r:
                print(r)


# === MAIN ===
if __name__ == "__main__":
    run(debug=True)   # 🔁 set False for multiprocessing